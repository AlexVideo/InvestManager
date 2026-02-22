# doc_generator.py
# -*- coding: utf-8 -*-
"""
Генератор черновика служебной записки по ревизии бюджета.

Поведение:
- Если установлен python-docx, создаём оформленный .docx c заголовками и абзацами.
- Если python-docx недоступен, собираем минимальный DOCX вручную, но с тем же текстом.

Интерфейс:
generate_revision_memo(src_project: str, dst_project: str, amount: float, date: str, note: str = "") -> str
Возвращает абсолютный путь к созданному .docx в папке data/memos/.
"""

from __future__ import annotations
import os
import zipfile
import textwrap
from xml.sax.saxutils import escape

import db
import utils  # используем форматирование сумм из utils.money

try:
    # Не критично, если модуля нет — есть резервная сборка DOCX
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_PYDOCX = True
except Exception:
    _HAS_PYDOCX = False


def _format_amount_tenge(amount: float) -> str:
    """Формат суммы в стиле интерфейса: 2 000 000,00 тенге без НДС."""
    return f"{utils.money(amount)} тенге без НДС"


def _year_from_iso(date_str: str) -> str:
    """Извлечь год из строки 'YYYY-MM-DD'. Если формат неожиданный — вернуть пусто."""
    if not date_str:
        return ""
    parts = date_str.split("-")
    return parts[0] if parts and len(parts[0]) == 4 else ""


def _safe_name(name: str) -> str:
    """Безопасное имя для файла."""
    return (name or "").replace("/", "-").replace("\\", "-")


def _build_text(src_project: str, dst_project: str, amount: float, date: str, note: str) -> str:
    """Собрать текст служебной записки (для fallback-режима и как источник контента)."""
    year = _year_from_iso(date) or "текущий"
    amount_txt = _format_amount_tenge(amount)
    note_txt = note.strip() if (note and note.strip()) else "—"

    body = textwrap.dedent(f"""
        Служебная записка
        Тема: Ревизия бюджетных средств

        В рамках реализации сводного плана инвестиций на {year} год по руднику «Жалпак» прошу рассмотреть возможность проведения ревизии бюджетных средств.

        Предлагается перераспределить средства:
        - Источник финансирования: «{src_project}»
        - Назначение: «{dst_project}»
        - Сумма: {amount_txt}
        - Дата проведения: {date}

        Обоснование:
        {note_txt}

        В связи с вышеизложенным, прошу Вас оказать содействие в проведении ревизии указанной статьи и поручить ДЗиМТО начать закупочные процедуры.
    """).strip()
    return body


# ===== Fallback: минимальный DOCX без python-docx =====

_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""

_RELS = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>
"""

_DOC_TPL = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:o="urn:schemas-microsoft-com:office:office"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
 xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
 xmlns:w10="urn:schemas-microsoft-com:office:word"
 xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
 xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
 xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
 xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
 xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" mc:Ignorable="w14 wp14">
  <w:body>
    {paras}
    <w:sectPr/>
  </w:body>
</w:document>
"""

def _xml_p(text: str) -> str:
    """Один параграф Word с экранированным текстом и переносами."""
    # Разобьём исходный текст на строки; пустые строки — это отдельные параграфы
    lines = text.splitlines()
    xml = []
    for ln in lines:
        if not ln.strip():
            # Пустая строка — просто пустой параграф
            xml.append("<w:p/>")
            continue
        xml.append(f"<w:p><w:r><w:t>{escape(ln)}</w:t></w:r></w:p>")
    return "\n".join(xml)

def _docx_minimal(text: str, out_path: str) -> None:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", _CONTENT_TYPES)
        z.writestr("_rels/.rels", _RELS)
        doc_xml = _DOC_TPL.format(paras=_xml_p(text))
        z.writestr("word/document.xml", doc_xml)


# ===== Основная функция =====

def generate_revision_memo(src_project: str, dst_project: str, amount: float, date: str, note: str = "", project_id: int | None = None) -> str:
    """
    Создаёт черновик служебной записки в папке проекта Files/{база}/{проект}/ и возвращает абсолютный путь к .docx.
    project_id — проект, в чью папку сохранять (карточка, из которой вызвали); если None — в Files/memos (совместимость).
    Имя файла: Мемо_ревизия_[Источник]_→_[Назначение]_[YYYY-MM-DD].docx
    """
    if project_id is not None:
        memos_dir = db.get_project_files_dir(project_id)
    else:
        db.ensure_data_dirs()
        memos_dir = os.path.join(db.FILES_DIR, "memos")
        os.makedirs(memos_dir, exist_ok=True)

    safe_src = _safe_name(src_project)
    safe_dst = _safe_name(dst_project)
    fname = f"Мемо_ревизия_{safe_src}_→_{safe_dst}_{date}.docx"
    out_path = os.path.join(memos_dir, fname)
    out_path = os.path.abspath(out_path)

    # Содержимое
    text = _build_text(src_project, dst_project, amount, date, note)

    # 1) Если есть python-docx — оформленный документ
    if _HAS_PYDOCX:
        doc = Document()

        # Заголовок
        p_title = doc.add_paragraph("Служебная записка")
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_title.runs[0]
        run.bold = True
        run.font.size = Pt(14)

        p_theme = doc.add_paragraph("Тема: Ревизия бюджетных средств")
        p_theme.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_theme.runs[0].font.size = Pt(11)

        doc.add_paragraph()  # пустая строка

        # Вводная
        year = _year_from_iso(date) or "текущий"
        p_intro = doc.add_paragraph(
            f"В рамках реализации сводного плана инвестиций на {year} год по руднику «Жалпак» "
            f"прошу рассмотреть возможность проведения ревизии бюджетных средств."
        )
        p_intro.runs[0].font.size = Pt(11)

        doc.add_paragraph()

        # Блок предложения
        doc.add_paragraph("Предлагается перераспределить средства:")
        ul = [
            f"Источник финансирования: «{src_project}»",
            f"Назначение: «{dst_project}»",
            f"Сумма: {_format_amount_tenge(amount)}",
            f"Дата проведения: {date}",
        ]
        for item in ul:
            para = doc.add_paragraph(item, style=None)
            para_format = para.paragraph_format
            para_format.left_indent = Pt(12)  # лёгкий отступ
            # маркер «- » спереди
            para.runs[0].text = "- " + para.runs[0].text
            para.runs[0].font.size = Pt(11)

        doc.add_paragraph()

        # Обоснование
        doc.add_paragraph("Обоснование:")
        p_note = doc.add_paragraph(note.strip() if note and note.strip() else "—")
        p_note.paragraph_format.left_indent = Pt(12)
        for r in p_note.runs:
            r.font.size = Pt(11)

        doc.add_paragraph()

        # Финальный абзац
        p_final = doc.add_paragraph(
            "В связи с вышеизложенным, прошу Вас оказать содействие в проведении ревизии указанной статьи "
            "и поручить ДЗиМТО начать закупочные процедуры."
        )
        p_final.runs[0].font.size = Pt(11)

        # Сохранение
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        doc.save(out_path)
        return out_path

    # 2) Fallback: минимальный DOCX без python-docx
    _docx_minimal(text, out_path)
    return out_path
